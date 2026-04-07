#!/usr/bin/env python3
"""
3GPP RAG V2 - Parse Specs with Version Support
Version extraction: 1) docx cover, 2) directory name, 3) default
"""

import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

# 导入配置加载器
try:
    from config_loader import load_config, get_path, ConfigError
except ImportError:
    sys.path.insert(0, str(Path(__file__).parent))
    from config_loader import load_config, get_path, ConfigError

# 延迟加载配置
_rag_config = None

def get_rag_config():
    """获取 RAG 配置（单例模式）"""
    global _rag_config
    if _rag_config is None:
        _rag_config = load_config()
    return _rag_config

def WORK_DIR() -> Path:
    """获取工作目录"""
    return get_path(get_rag_config(), "work_dir")

def SPEC_BASE_DIR() -> Path:
    """获取协议根目录"""
    return get_path(get_rag_config(), "protocol_base")

def LOG_DIR() -> Path:
    """获取日志目录"""
    return get_path(get_rag_config(), "log_dir")

def LOG_FILE() -> Path:
    """获取日志文件路径"""
    return LOG_DIR() / "rag_parse_specs.log"

DEFAULT_VERSION = "Rel-19"

# Protocol list with version
PROTOCOLS = [
    ("Rel-19", "38_series", "38300-j20.zip"),
    ("Rel-19", "38_series", "38321-j20.zip"),
    ("Rel-19", "38_series", "38322-j20.zip"),
    ("Rel-19", "38_series", "38331-j20.zip"),
    ("Rel-19", "38_series", "38101-1-j50.zip"),
    ("Rel-19", "38_series", "38101-2-j40.zip"),
    ("Rel-19", "38_series", "38101-3-j50.zip"),
    ("Rel-19", "38_series", "38101-4-j20.zip"),
    ("Rel-19", "38_series", "38101-5-j40.zip"),
    ("Rel-19", "38_series", "38104-j40.zip"),
    # New PHY protocols
    ("Rel-19", "38_series", "38211-j30.zip"),
    ("Rel-19", "38_series", "38212-j30.zip"),
    ("Rel-19", "38_series", "38213-j30.zip"),
    ("Rel-19", "38_series", "38214-j30.zip"),
    # 38.133 (RRM) - may timeout
    ("Rel-19", "38_series", "38133-j40.zip"),
]

def log(msg):
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{ts}] {msg}"
    print(line)
    log_path = LOG_FILE()
    log_path.parent.mkdir(parents=True, exist_ok=True)
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(line + "\n")


def extract_version_from_docx(docx_path):
    """Extract version from docx cover page."""
    try:
        doc = Document(str(docx_path))
        for para in doc.paragraphs[:50]:
            text = para.text.lower()
            if "release 19" in text or "rel-19" in text or "r19" in text:
                return "Rel-19"
            if "release 20" in text or "rel-20" in text or "r20" in text:
                return "Rel-20"
    except Exception as e:
        log(f"  Warning: Could not extract version from docx: {e}")
    return None


def get_version(zip_path, docx_path, default_version):
    """Get version: docx -> path -> default"""
    version = extract_version_from_docx(docx_path)
    if version:
        return version
    # Check path for version folder
    for part in zip_path.parts:
        if part.startswith("Rel-"):
            return part
    return default_version


def parse_clause_number(text):
    match = re.match(r'^(\d+(?:\.\d+)*)\s+', text.strip())
    return match.group(1) if match else None


def get_clause_level(clause_number):
    return len(clause_number.split('.'))


def parse_docx(docx_path, version):
    """Parse a docx file."""
    log(f"  Parsing: {docx_path.name}")
    try:
        doc = Document(str(docx_path))
    except Exception as e:
        log(f"  ERROR: {e}")
        return None
    
    match = re.match(r'(\d{2})(\d{3})', docx_path.stem.split('-')[0])
    spec_id = f"{match.group(1)}.{match.group(2)}" if match else "unknown"
    
    clauses = []
    current_clause = None
    current_content = []
    
    def save_clause():
        if current_clause:
            cn, title = current_clause
            content = '\n'.join(current_content).strip()
            clauses.append({
                "clause_number": cn,
                "title": title,
                "content": content[:5000],
                "level": get_clause_level(cn),
                "version": version
            })
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        is_heading = style.startswith("Heading") or style.startswith("标题")
        clause_num = parse_clause_number(text)
        
        if is_heading and clause_num:
            save_clause()
            title = re.sub(r'^\d+(?:\.\d+)*\s*', '', text).strip()
            current_clause = (clause_num, title)
            current_content = [f"## {clause_num} {title}"]
        elif current_clause and text:
            current_content.append(text)
    
    save_clause()
    return {"spec_id": spec_id, "version": version, "clauses_count": len(clauses), "clauses": clauses}


def process_spec(version, series, zip_name):
    """Process a single spec."""
    zip_path = SPEC_BASE_DIR / version / series / zip_name
    extract_dir = WORK_DIR / "extracted_v2" / zip_name.replace('.zip', '')
    
    log(f"\nProcessing: {zip_name}")
    log(f"  Path: {zip_path}")
    
    if not zip_path.exists():
        log(f"  ERROR: Zip not found")
        return None
    
    size_mb = zip_path.stat().st_size / (1024 * 1024)
    log(f"  Size: {size_mb:.2f} MB")
    
    # Extract
    extract_dir.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(zip_path, 'r') as zf:
            zf.extractall(extract_dir)
    except Exception as e:
        log(f"  ERROR extracting: {e}")
        return None
    
    docx_files = list(extract_dir.glob("*.docx"))
    log(f"  Found {len(docx_files)} docx files")
    
    all_clauses = []
    spec_id = None
    
    for docx_file in docx_files:
        actual_version = get_version(zip_path, docx_file, version)
        result = parse_docx(docx_file, actual_version)
        if result:
            if spec_id is None:
                spec_id = result["spec_id"]
            all_clauses.extend(result["clauses"])
            log(f"    {docx_file.name}: {result['clauses_count']} clauses, version={actual_version}")
    
    return {
        "spec_name": spec_id or zip_name,
        "zip_name": zip_name,
        "version": version,
        "clauses": all_clauses,
        "clause_count": len(all_clauses)
    }


def main():
    log("=" * 60)
    log("RAG V2 - Parse Specs with Version Support")
    log("=" * 60)
    
    (WORK_DIR / "extracted_v2").mkdir(parents=True, exist_ok=True)
    
    all_specs = []
    total_clauses = 0
    failed_specs = []
    
    for version, series, zip_name in PROTOCOLS:
        try:
            result = process_spec(version, series, zip_name)
            if result and result["clause_count"] > 0:
                all_specs.append(result)
                total_clauses += result["clause_count"]
                log(f"  SUCCESS: {result['spec_name']} ({result['version']}) - {result['clause_count']} clauses")
            else:
                failed_specs.append(zip_name)
                log(f"  FAILED: {zip_name}")
        except Exception as e:
            log(f"  EXCEPTION: {zip_name} - {e}")
            failed_specs.append(zip_name)
    
    # Save results
    output_file = WORK_DIR / "all_specs_v2.json"
    data = {
        "status": "ok",
        "total_specs": len(all_specs),
        "total_clauses": total_clauses,
        "specs": all_specs,
        "failed_specs": failed_specs,
        "timestamp": datetime.now().isoformat()
    }
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    log("\n" + "=" * 60)
    log(f"Parse Complete!")
    log(f"  Total specs: {len(all_specs)}")
    log(f"  Total clauses: {total_clauses}")
    log(f"  Failed: {len(failed_specs)}")
    if failed_specs:
        log(f"  Failed: {failed_specs}")
    log("=" * 60)


if __name__ == "__main__":
    main()
