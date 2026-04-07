#!/usr/bin/env python3
"""
3GPP RAG V3 - Batch add all missing 38-series specs
Processes by size (small first, large last), with timeout and retry.
"""

import sys
import json
import re
import zipfile
import traceback
from pathlib import Path
from datetime import datetime
import multiprocessing

from docx import Document
import chromadb
from chromadb.utils import embedding_functions

# 导入配置加载器
try:
    from config_loader import load_config, get_path, ConfigError
except ImportError:
    sys.path.insert(0, str(Path(__file__).parent))
    from config_loader import load_config, get_path, ConfigError

# 延迟加载配置
_rag_config = None

def get_rag_config():
    """获取 RAG 配置（单例模式）"""
    global _rag_config
    if _rag_config is None:
        _rag_config = load_config()
    return _rag_config

def WORK_DIR() -> Path:
    """获取工作目录"""
    return get_path(get_rag_config(), "work_dir")

def PROTOCOL_BASE() -> Path:
    """获取协议根目录"""
    return get_path(get_rag_config(), "protocol_base")

def LOG_DIR() -> Path:
    """获取日志目录"""
    return get_path(get_rag_config(), "log_dir")

def DB_DIR() -> Path:
    """获取数据库目录（旧版兼容）"""
    return WORK_DIR() / "chroma_db_complete_v2"

def LOG_FILE() -> Path:
    """获取日志文件路径"""
    return LOG_DIR() / "rag_batch_add.log"

DEFAULT_VERSION = "Rel-19"
# Timeout per spec in seconds
TIMEOUT_SMALL = 120    # < 5MB
TIMEOUT_MED = 180      # 5-15MB
TIMEOUT_LARGE = 300    # 15-35MB
TIMEOUT_XL = 600       # > 35MB


def log(msg):
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{ts}] {msg}"
    print(line)
    log_path = LOG_FILE()
    log_path.parent.mkdir(parents=True, exist_ok=True)
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(line + "\n")


def get_zip_size_mb(spec_name):
    """Get zip file size in MB for a spec."""
    base = PROTOCOL_BASE() / DEFAULT_VERSION / "38_series"
    # Try to find matching zip
    pattern = spec_name.replace(".", "")
    for zf in sorted(base.glob(f"{pattern}*.zip")):
        return zf.stat().st_size / (1024 * 1024)
    return 0


def find_all_missing_specs():
    """Find all specs not yet in database."""
    client = chromadb.PersistentClient(path=str(DB_DIR))
    collection = client.get_collection(name="3gpp_complete")
    results = collection.get()
    existing = set(m['spec'] for m in results['metadatas'])

    base = PROTOCOL_BASE / DEFAULT_VERSION / "38_series"
    all_zips = list(base.glob("*.zip"))

    missing = []
    seen_specs = set()
    for zf in all_zips:
        name = zf.stem  # e.g. "38133-j40"
        match = re.match(r'(\d{2})(\d{3})', name.split('-')[0])
        if not match:
            continue
        spec_id = f"{match.group(1)}.{match.group(2)}"
        # Handle sub-versions like 38.101-1
        sub_match = re.match(r'(\d{2})(\d{3})-(\d+)', name.split('-')[0])
        if sub_match:
            spec_id = f"{sub_match.group(1)}.{sub_match.group(2)}-{sub_match.group(3)}"

        # For 38.101 series, use base spec number
        base_spec = spec_id.split('-')[0] if '-' in spec_id else spec_id

        # Skip 38.101 sub-files already covered
        if base_spec == "38.101" and spec_id != "38.101":
            continue

        if spec_id not in existing and spec_id not in seen_specs:
            size_mb = zf.stat().st_size / (1024 * 1024)
            missing.append({
                "spec": spec_id,
                "zip_name": name,
                "size_mb": round(size_mb, 2),
                "timeout": _get_timeout(size_mb),
            })
            seen_specs.add(spec_id)

    # Sort by size ascending
    missing.sort(key=lambda x: x["size_mb"])
    return missing


def _get_timeout(size_mb):
    if size_mb < 5:
        return TIMEOUT_SMALL
    elif size_mb < 15:
        return TIMEOUT_MED
    elif size_mb < 35:
        return TIMEOUT_LARGE
    else:
        return TIMEOUT_XL


def parse_clause_number(text):
    match = re.match(r'^(\d+(?:\.\d+)*)\s+', text.strip())
    return match.group(1) if match else None


def extract_version_from_docx(docx_path):
    try:
        doc = Document(str(docx_path))
        for para in doc.paragraphs[:50]:
            text = para.text.lower()
            if "release 19" in text or "rel-19" in text:
                return "Rel-19"
            if "release 20" in text or "rel-20" in text:
                return "Rel-20"
    except:
        pass
    return None


def add_single_spec(spec_info):
    """Add a single spec to database. Returns (spec, clause_count, status, error_msg)."""
    spec = spec_info["spec"]
    zip_name = spec_info["zip_name"]

    base = PROTOCOL_BASE / DEFAULT_VERSION / "38_series"
    zip_path = base / f"{zip_name}.zip"

    if not zip_path.exists():
        return (spec, 0, "SKIP", f"Zip not found: {zip_path}")

    import tempfile
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            with zipfile.ZipFile(str(zip_path), 'r') as z:
                z.extractall(tmpdir)

            docx_files = list(Path(tmpdir).glob("*.docx"))
            if not docx_files:
                return (spec, 0, "SKIP", "No docx files")

            # Extract version
            version = DEFAULT_VERSION
            for df in docx_files:
                v = extract_version_from_docx(str(df))
                if v:
                    version = v
                    break

            # Parse all docx
            all_clauses = []
            for df in docx_files:
                try:
                    doc = Document(str(df))
                except:
                    continue

                current_clause = None
                current_content = []

                def save_clause():
                    if current_clause:
                        cn, title = current_clause
                        content = '\n'.join(current_content).strip()
                        if len(content) >= 10:
                            all_clauses.append({
                                "clause_number": cn,
                                "title": title,
                                "content": content[:5000],
                                "level": len(cn.split('.')),
                                "version": version,
                            })

                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style = para.style.name if para.style else ""
                    is_heading = style.startswith("Heading") or style.startswith("\u6807\u9898")
                    clause_num = parse_clause_number(text)

                    if is_heading and clause_num:
                        save_clause()
                        title = re.sub(r'^\d+(?:\.\d+)*\s*', '', text).strip()
                        current_clause = (clause_num, title)
                        current_content = [f"## {clause_num} {title}"]
                    elif current_clause and text:
                        current_content.append(text)

                save_clause()

            if not all_clauses:
                return (spec, 0, "SKIP", "No clauses parsed")

            # Add to database
            client = chromadb.PersistentClient(path=str(DB_DIR))
            collection = client.get_collection(name="3gpp_complete")

            # Remove existing
            existing = collection.get(where={"spec": spec})
            if existing['ids']:
                collection.delete(ids=existing['ids'])

            # Batch add
            batch_ids = []
            batch_docs = []
            batch_metas = []

            for clause in all_clauses:
                clause_num = clause['clause_number']
                title = clause['title']
                content = clause['content']
                doc_id = f"{spec}_{clause_num}"

                embed_text = (
                    f"Specification: {spec}\n"
                    f"Version: {version}\n"
                    f"Clause: {clause_num}\n"
                    f"Title: {title}\n\n"
                    f"Content:\n{content[:2000]}"
                )

                batch_ids.append(doc_id)
                batch_docs.append(embed_text)
                batch_metas.append({
                    "spec": spec,
                    "version": version,
                    "clause": clause_num,
                    "title": title,
                    "level": clause['level'],
                })

                if len(batch_ids) >= 100:
                    collection.add(ids=batch_ids, documents=batch_docs, metadatas=batch_metas)
                    batch_ids = []
                    batch_docs = []
                    batch_metas = []

            if batch_ids:
                collection.add(ids=batch_ids, documents=batch_docs, metadatas=batch_metas)

            return (spec, len(all_clauses), "OK", version)

    except Exception as e:
        return (spec, 0, "ERROR", str(e)[:200])


def run_with_timeout(spec_info, timeout):
    """Run add_single_spec with timeout."""
    pool = multiprocessing.Pool(1)
    try:
        result = pool.apply_async(add_single_spec, (spec_info,))
        return result.get(timeout=timeout)
    except multiprocessing.TimeoutError:
        pool.terminate()
        pool.join()
        return (spec_info["spec"], 0, "TIMEOUT", f"Exceeded {timeout}s")
    except Exception as e:
        pool.terminate()
        pool.join()
        return (spec_info["spec"], 0, "ERROR", str(e)[:200])
    finally:
        try:
            pool.close()
        except:
            pass


def main():
    if sys.platform == "win32":
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    log("=" * 70)
    log("BATCH ADD - All Missing 38-Series Specs")
    log("=" * 70)

    # Find missing specs
    missing = find_all_missing_specs()
    log(f"Found {len(missing)} missing specs")

    if not missing:
        log("Nothing to add!")
        return

    log("")
    for s in missing:
        log(f"  {s['spec']} ({s['zip_name']}) - {s['size_mb']} MB - timeout {s['timeout']}s")

    # Process in batches
    success = []
    failed = []
    skipped = []

    log("")
    log("=" * 70)
    log("Phase 1: Small and Medium files")
    log("=" * 70)

    for spec_info in missing:
        if spec_info["size_mb"] >= 15:
            continue  # Large files processed in Phase 2

        spec = spec_info["spec"]
        timeout = spec_info["timeout"]
        size = spec_info["size_mb"]
        log(f"\nProcessing: {spec} ({size} MB, timeout {timeout}s)")

        result = run_with_timeout(spec_info, timeout)
        spec, count, status, msg = result

        if status == "OK":
            log(f"  SUCCESS: {spec} - {count} clauses ({msg})")
            success.append(spec_info)
        elif status == "SKIP":
            log(f"  SKIP: {spec} - {msg}")
            skipped.append(spec_info)
        else:
            log(f"  {status}: {spec} - {msg}")
            failed.append(spec_info)

    log("")
    log("=" * 70)
    log("Phase 2: Large files (>=15MB)")
    log("=" * 70)

    large_files = [s for s in missing if s["size_mb"] >= 15]
    for spec_info in large_files:
        spec = spec_info["spec"]
        timeout = spec_info["timeout"]
        size = spec_info["size_mb"]
        log(f"\nProcessing: {spec} ({size} MB, timeout {timeout}s)")

        result = run_with_timeout(spec_info, timeout)
        spec, count, status, msg = result

        if status == "OK":
            log(f"  SUCCESS: {spec} - {count} clauses ({msg})")
            success.append(spec_info)
        elif status == "SKIP":
            log(f"  SKIP: {spec} - {msg}")
            skipped.append(spec_info)
        else:
            log(f"  {status}: {spec} - {msg}")
            failed.append(spec_info)

    # Phase 3: Retry failed with extended timeout
    if failed:
        log("")
        log("=" * 70)
        log("Phase 3: Retry failed specs with extended timeout (900s)")
        log("=" * 70)

        retry_list = list(failed)
        failed = []

        for spec_info in retry_list:
            spec = spec_info["spec"]
            log(f"\nRetrying: {spec} (timeout 900s)")

            result = run_with_timeout(spec_info, 900)
            spec, count, status, msg = result

            if status == "OK":
                log(f"  SUCCESS: {spec} - {count} clauses ({msg})")
                success.append(spec_info)
            else:
                log(f"  STILL {status}: {spec} - {msg}")
                failed.append(spec_info)

    # Phase 4: Try sub-file approach for remaining failures
    if failed:
        log("")
        log("=" * 70)
        log("Phase 4: Attempting sub-file approach for remaining failures")
        log("=" * 70)

        for spec_info in failed:
            spec = spec_info["spec"]
            log(f"\nAttempting sub-file approach: {spec}")

            try:
                base = PROTOCOL_BASE / DEFAULT_VERSION / "38_series"
                zip_path = base / f"{spec_info['zip_name']}.zip"
                import tempfile
                with tempfile.TemporaryDirectory() as tmpdir:
                    with zipfile.ZipFile(str(zip_path), 'r') as z:
                        z.extractall(tmpdir)

                    docx_files = sorted(Path(tmpdir).glob("*.docx"))
                    log(f"  {len(docx_files)} docx files found")

                    added_total = 0
                    added_files = 0
                    for df in docx_files:
                        try:
                            mini_info = {
                                "spec": spec,
                                "zip_name": spec_info["zip_name"],
                                "size_mb": spec_info["size_mb"],
                                "timeout": 120,
                            }
                            # Process each file individually
                            doc = Document(str(df))

                            version = DEFAULT_VERSION
                            v = extract_version_from_docx(str(df))
                            if v:
                                version = v

                            clauses = []
                            current_clause = None
                            current_content = []

                            def save():
                                if current_clause:
                                    cn, title = current_clause
                                    content = '\n'.join(current_content).strip()
                                    if len(content) >= 10:
                                        clauses.append({
                                            "clause_number": cn,
                                            "title": title,
                                            "content": content[:5000],
                                            "level": len(cn.split('.')),
                                            "version": version,
                                        })

                            for para in doc.paragraphs:
                                text = para.text.strip()
                                if not text:
                                    continue
                                style = para.style.name if para.style else ""
                                is_heading = style.startswith("Heading") or style.startswith("\u6807\u9898")
                                clause_num = parse_clause_number(text)
                                if is_heading and clause_num:
                                    save()
                                    title = re.sub(r'^\d+(?:\.\d+)*\s*', '', text).strip()
                                    current_clause = (clause_num, title)
                                    current_content = [f"## {clause_num} {title}"]
                                elif current_clause and text:
                                    current_content.append(text)
                            save()

                            if clauses:
                                client = chromadb.PersistentClient(path=str(DB_DIR))
                                collection = client.get_collection(name="3gpp_complete")
                                batch_ids = []
                                batch_docs = []
                                batch_metas = []
                                for c in clauses:
                                    doc_id = f"{spec}_{c['clause_number']}"
                                    embed_text = f"Specification: {spec}\nVersion: {version}\nClause: {c['clause_number']}\nTitle: {c['title']}\n\nContent:\n{c['content'][:2000]}"
                                    batch_ids.append(doc_id)
                                    batch_docs.append(embed_text)
                                    batch_metas.append({"spec": spec, "version": version, "clause": c['clause_number'], "title": c['title'], "level": c['level']})
                                if batch_ids:
                                    collection.add(ids=batch_ids, documents=batch_docs, metadatas=batch_metas)
                                added_total += len(clauses)
                                added_files += 1
                                log(f"    {df.name}: {len(clauses)} clauses")
                        except Exception as e:
                            log(f"    {df.name}: ERROR - {str(e)[:100]}")

                    if added_total > 0:
                        log(f"  PARTIAL SUCCESS: {spec} - {added_total} clauses from {added_files}/{len(docx_files)} files")
                        success.append(spec_info)
                        if spec_info in failed:
                            failed.remove(spec_info)
                    else:
                        log(f"  FAILED: {spec} - no clauses from any file")

            except Exception as e:
                log(f"  FAILED: {spec} - {str(e)[:200]}")

    # Final summary
    log("")
    log("=" * 70)
    log("FINAL SUMMARY")
    log("=" * 70)
    log(f"  Success: {len(success)}")
    log(f"  Skipped:  {len(skipped)}")
    log(f"  Failed:   {len(failed)}")

    if success:
        log("")
        log("  Added specs:")
        for s in success:
            log(f"    {s['spec']} ({s['size_mb']} MB)")

    if failed:
        log("")
        log("  Failed specs (manual intervention needed):")
        for s in failed:
            log(f"    {s['spec']} ({s['size_mb']} MB) - {s['zip_name']}")

    # Database total
    try:
        client = chromadb.PersistentClient(path=str(DB_DIR))
        collection = client.get_collection(name="3gpp_complete")
        log(f"\n  Database total: {collection.count()} documents")
    except:
        pass

    log("=" * 70)


if __name__ == "__main__":
    main()
